
import streamlit as st
import pdfplumber
import os
from docx import Document
from io import BytesIO
from openai import OpenAI

# API-Key laden (z. B. über Secrets oder Umgebungsvariable)
api_key = st.secrets.get("OPENAI_API_KEY") or os.getenv("OPENAI_API_KEY")

if not api_key:
    st.error("❌ Kein OpenAI API-Key gefunden. Bitte in secrets oder Umgebungsvariablen eintragen.")
    st.stop()

client = OpenAI(api_key=api_key)

st.set_page_config(page_title="Makro-Ausblick Generator", layout="wide")
st.title("📊 Volkswirtschaftlicher Ausblick – Mittelfristplanung Bank")

st.markdown("Lade hier relevante PDF-Dokumente hoch (z. B. Bundesbankberichte, EZB-Ausblicke, Research-Material):")

uploaded_files = st.file_uploader("PDF-Dateien auswählen", type=["pdf"], accept_multiple_files=True)

if uploaded_files:
    full_text = ""
    for file in uploaded_files:
        with pdfplumber.open(file) as pdf:
            for page in pdf.pages:
                full_text += page.extract_text() + "\n"

    if not full_text.strip():
        st.error("❌ Konnte keinen Text aus den PDFs extrahieren.")
        st.stop()

    st.success("✅ PDF-Dateien erfolgreich gelesen.")

    st.markdown("### 🔍 Generiere volkswirtschaftlichen Ausblick...")

    with st.spinner("KI analysiert die Daten und erstellt den Ausblick..."):
        prompt = f"""
        Du bist ein volkswirtschaftlicher Analyst einer Bank.
        Erstelle einen **detaillierten volkswirtschaftlichen Ausblick** für die nächsten 3–5 Jahre
        für die **Mittelfristplanung einer Bank** auf Basis folgender Texte:

        ---
        {full_text[:15000]}
        ---

        Berücksichtige:
        - BIP-Wachstum (global, EU, Deutschland)
        - Inflationstrends und -risiken
        - Leitzinsprognosen (EZB, FED)
        - Arbeitsmarkt
        - Immobilien- und Finanzmärkte
        - Politische/geopolitische Risiken
        - Risiken & Chancen für Banken

        Struktur:
        1. Makroökonomisches Umfeld
        2. Kapitalmärkte & Zinsen
        3. Auswirkungen auf Banken
        4. Relevante Annahmen für Mittelfristplanung
        5. Fazit / Handlungsempfehlungen
        """

        response = client.chat.completions.create(
            model="gpt-4",
            messages=[{"role": "user", "content": prompt}],
            temperature=0.3,
        )

        outlook = response.choices[0].message.content

        st.success("✅ Ausblick erstellt.")
        st.markdown("### 📄 Vorschau")
        st.markdown(outlook)

        doc = Document()
        doc.add_heading("Volkswirtschaftlicher Ausblick – Mittelfristplanung", 0)
        for para in outlook.split("\n"):
            doc.add_paragraph(para.strip())

        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        st.download_button(
            label="📥 Word-Dokument herunterladen",
            data=buffer,
            file_name="Volkswirtschaftlicher_Ausblick.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
